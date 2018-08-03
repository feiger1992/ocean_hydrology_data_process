from PyQt5.QtCore import pyqtSignal, QObject,QThread
from PyQt5.QtGui import QDoubleValidator
from PyQt5.QtWidgets import QApplication, QMainWindow, QAction, QLabel, QVBoxLayout, QFormLayout, QLineEdit, QSpinBox, \
     QFileDialog, QTextBrowser, QGroupBox, QHBoxLayout, QTabWidget, QPushButton, QWidget, QCheckBox, \
    QMessageBox, QDialog, QDateTimeEdit, QComboBox,  QErrorMessage, QDoubleSpinBox, QInputDialog, \
    QTextEdit,QTableWidget,QTableWidgetItem

from threading import Thread
from  docx import Document
from  docx.shared import Pt
from  docx.oxml.ns import qn
from docx.shared import RGBColor
import sys, time
import pandas


# 从__main__执行时使用以下导入
from Metocean.tide import Process_Tide
from Metocean.current import Single_Tide_Point, Current_pre_process,Read_Report

# 从本文件执行时使用以下导入

#

from dxfwrite import DXFEngine as dxf

import numpy as np
import matplotlib
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.backends.backend_qt5agg import NavigationToolbar2QT as Navigation


dir_of_file = lambda filename: filename.replace(list(filename.split(sep='\\'))[-1], '')
str_from_df = lambda item: item.values[0]
num_from_df = lambda item: str(round(item.values[0], 2))


# 通知栏         QMessageBox.information(self,'通知标题',"内容")

class M_window(QMainWindow):
    def __init__(self):
        super().__init__()
        self.show_content = {}
        self.initUI()
        global sig
        sig = All_Signal_event()
        self.set_signal_event()

    def show_pic(self, fig):

        self.diag = QWidget()
        canvas = FigureCanvas(fig)
        canvas.setParent(self.diag)
        self.mpl_toolbar = Navigation(canvas, self.diag)

        vbox = QVBoxLayout()
        vbox.addWidget(self.mpl_toolbar)
        vbox.addWidget(canvas)
        self.diag.setLayout(vbox)

        # self.shower.addTab(self.canvas,'HAHAHA')
        self.win = QMainWindow()
        self.win.setCentralWidget(self.diag)
        self.win.show()


    def initUI(self):

        self.setWindowTitle('中交三航院水文数据处理软件')
        self.resize(1000, 618)
        self.statusBar()
        # self.createMenu()
        centralWidget = QWidget()

        self.shower = QTabWidget()
        self.gen_welcome_web()
        self.shower.addTab(self.welcome_tab, '欢迎页面')
        self.allGroup = QGroupBox()
        self.mid_l_Group = QGroupBox()
        self.mid_s_Group = QGroupBox("测验概况")
        self.Tabs = QTabWidget()
        self.Tabs.setMaximumHeight(350)
        self.Generate_all_button = QPushButton("生成报告及文档")
        centralLayout = QVBoxLayout()
        centralLayout.addWidget(self.shower, stretch=40)
        centralLayout.addWidget(self.allGroup, stretch=10)
        centralWidget.setLayout(centralLayout)
        self.setCentralWidget(centralWidget)

        allGroupLayout = QHBoxLayout()
        # allGroupLayout.addWidget(self.mid_l_Group, stretch=3)
        allGroupLayout.addWidget(self.Tabs, stretch=30)
        self.allGroup.setLayout(allGroupLayout)
        # 测试潮流，将潮位部分取消
        self.tides_control = TideTab()
        self.Tabs.addTab(self.tides_control, "潮位")

        self.current_control = Current_Tab()
        self.Tabs.addTab(self.current_control, '潮流')

        self.wind_control = Wind_Tab()
        self.Tabs.addTab(self.wind_control, '气象')

        self.wave_control = Wave_tab()
        self.Tabs.addTab(self.wave_control, '波浪')

        mid_l_GroupLayout = QVBoxLayout()
        mid_l_GroupLayout.addWidget(self.mid_s_Group)
        mid_l_GroupLayout.addWidget(self.Generate_all_button)
        self.mid_l_Group.setLayout(mid_l_GroupLayout)

        mid_s_GroupLayout = QFormLayout()
        Project_Name = QLineEdit()
        Sea_Area = QLineEdit()
        Button_Save_Info = QPushButton("确定")
        mid_s_GroupLayout.addRow(QLabel("工程名称"), Project_Name)
        mid_s_GroupLayout.addRow(QLabel("测验海域"), Sea_Area)
        mid_s_GroupLayout.addRow(Button_Save_Info)
        self.mid_s_Group.setLayout(mid_s_GroupLayout)

    def gen_welcome_web(self):
        self.welcome_tab = QTextBrowser()
        self.welcome_tab.setHtml(
            r"<!DOCTYPE html>\n<html>\n<head>\n<title>Page Title</title>\n</head>\n<body>\n\n<h1>欢迎使用本水文处理软件，请在底部标签页选择需要使用的模块</h1>\n<p>潮位</p>\n<p>潮流</p>\n<p>气象</p>\n<p>波浪</p>\n<p>更多功能正在开发中···</p>\n\n</body>\n</html>\n".replace(
                '\\n', '')
        )

    def set_signal_event(self):
        sig.load_tide_file_done.connect(self.tide_site_element_event)
        sig.msg_to_show.connect(self.show_msg_statusbar)
        sig.load_current_file_done.connect(self.current_site_element_event)

        self.current_control.sig.show_all_out_signal.connect(self.display)

    def show_msg_statusbar(self, msg):
        self.statusBar().showMessage(msg)

    def createMenu(self):
        tide_load = QAction('载入数据', self)
        tide_load.setStatusTip('从Excel文件读入潮位数据，不同站点位于不同sheet,列标题应为Tide,Time')
        #        tide_load.triggered.connect(self.tide)

        generate_tide_sheet = QAction('&保存潮位报表', self)
        generate_tide_sheet.setStatusTip('保存潮位月报表')
        #       generate_tide_sheet.triggered.connect(self.save_tide)

        tide_statics = QAction('&各月潮位特征值汇总', self)
        tide_statics.setStatusTip('统计各月潮位极值,涨落潮历时等')
        #       tide_statics.triggered.connect(self.out_put_data)

        tide_harmonic_analysis = QAction('&调和分析', self)
        tide_harmonic_analysis.setStatusTip('对潮位数据进行调和分析，计算各分潮特征值')
        #      tide_harmonic_analysis.triggered.connect(self.harmonic_analysis)

        load_stream_E_N = QAction('&载入数据(格式为E-N)', self)
        load_stream_E_N.setStatusTip('对应sheet名应为‘E’,‘N’')

        load_stream_V_D = QAction('&载入数据(格式为V-D)', self)
        load_stream_V_D.setStatusTip('对应sheet名应为‘V’,‘D’')

        get_button = QAction('&去除底部干扰数据', self)
        get_button.setStatusTip('找底')

        constant_3 = QAction('&常数法（三点）', self)
        constant_6 = QAction('&常数法（六点）', self)

        power_3 = QAction('&幂函数法（三点）', self)
        power_6 = QAction('&幂函数法（六点）', self)

        stream_statics = QAction('&统计潮流特征', self)
        stream_statics.setStatusTip('统计最大涨落潮流速、涨落潮历时等特征数据并输出Excel')

        generate_stream_sheet = QAction('&生成潮流报表', self)

        menubar = self.menuBar()
        TIDE = menubar.addMenu('&潮位')
        TIDE.addAction(tide_load)

        TIDE.addAction(generate_tide_sheet)
        TIDE.addAction(tide_statics)
        TIDE.addAction(tide_harmonic_analysis)

        STREAM = menubar.addMenu('&潮流')
        LOAD_STREAM = STREAM.addMenu('&载入潮流数据')
        LOAD_STREAM.addAction(load_stream_E_N)
        LOAD_STREAM.addAction(load_stream_V_D)

        STREAM.addSeparator()

        STREAM.addAction(get_button)

        STREAM.addSeparator()

        FENCENG = STREAM.addMenu('&分层')
        FENCENG.addAction(constant_6)
        FENCENG.addAction(power_6)
        FENCENG.addSeparator()
        FENCENG.addAction(constant_3)
        FENCENG.addAction(power_3)

        STREAM.addSeparator()
        STREAM.addAction(stream_statics)
        STREAM.addAction(generate_stream_sheet)
        # 整体布局

    # def createTideTab(self):
    #    self.tideTab = QTabWidget

    def display(self, category, title, content, tip=None):
        if category == 'HTML':
            #Qweb = QWebEngineView()
            #Qweb.setHtml(content)
            #Qweb.show()
            #self.show_content.update({title + '(web)': Qweb})
            #self.shower.addTab(self.show_content[title + '(web)'], title + '(web)')

            #self.shower.setCurrentWidget(self.show_content[title + '(web)'])
            #if tip:
                #self.shower.setTabToolTip(self.shower.count() - 1, tip)
            pass
        if category == 'TXT':
            self.show_content.update({title: QTextEdit()})
            self.shower.addTab(self.show_content[title], title)
            # self.show_content[title].setHtml(content)#试试效果
            self.show_content[title].setReadOnly(True)
            self.show_content[title].setPlainText(content)
            self.shower.setCurrentWidget(self.show_content[title])
            if tip:
                self.shower.setTabToolTip(self.shower.count() - 1, tip)

    def display_tide_HTML(self, filename, site, tip):
        self.display(category='HTML', title=site, content=t[filename].html[site], tip=tip)

    def display_tide_TXT(self, filename, site, tip):
        columns = ['tide', 'tide_init', 'if_min', 'if_max', 'diff', 'raising_time', 'ebb_time']
        self.display(category='TXT', title=site,
                     content=t[filename].data[site].to_string(columns=columns, index_names=False).replace('None',
                                                                                                          ' ' * 4).replace(
                         'False', ' ' * 5).replace('NaN', ' ' * 3), tip=tip)

    def display_tide_harmonic_result(self, filename, site, tip):
        self.display(category='TXT', title=site + '调和分析结果', content=t[filename].harmonic_result[site].classic_style(),
                     tip=tip)

    def tide_process(self, filename, sitename, threshold):
        preprocess_thread = Thread(target=t[filename].preprocess, kwargs={'s': sitename, 'threshold': threshold})
        preprocess_thread.start()
        second = 0
        while preprocess_thread.is_alive():
            sig.msg_to_show.emit('正在对' + sitename + '站点数据进行预处理，用时' + str(round(second, 1)) + '秒')
            time.sleep(0.1)
            second += 0.1
        sig.msg_to_show.emit(sitename + '站数据预处理结束，请继续操作')

    def tide_site_element_event(self, file, site):
        self.tides_control.sites_control[site].signal.tide_preprocess.connect(self.tide_process)
        self.tides_control.sites_control[site].signal.show_tide_HTML_signal.connect(self.display_tide_HTML)
        self.tides_control.sites_control[site].signal.show_tide_TXT_signal.connect(self.display_tide_TXT)
        self.tides_control.sites_control[site].signal.show_tide_harmonic_result_signal.connect(
            self.display_tide_harmonic_result)
        self.tides_control.sites_control[site].signal.show_diff_signal.connect(self.show_pic)
        self.tides_control.sites_control[site].signal.show_tide_statics_signal.connect(self.display)

    def current_site_element_event(self, Point, tide_type):
        self.current_control.Point_control[Point].tides_control[tide_type].signal.show_current_arrow_signal.connect(
            self.show_pic)
        self.current_control.Point_control[Point].tides_control[
            tide_type].signal.show_current_proper_TXT_signal.connect(self.display)


class TideTab(QWidget):
    def __init__(self, parent=None):
        super(TideTab, self).__init__(parent)
        global t
        t = {}
        self.layout_element()

    def layout_element(self):

        self.tide_l_Group = QGroupBox("潮位文件")
        self.tide_file_button = QPushButton("导入潮位文件", clicked=self.open_file)
        self.auto_Preprocess = QCheckBox("打开文件时进行预处理")
        self.if_one_sheet = QCheckBox("仅载入首个工作表")
        self.check_multifile = QCheckBox("多个文件")
        self.check_multifile.setChecked(True)
        self.auto_Preprocess.setChecked(True)
        self.if_one_sheet.setChecked(True)
        # self.check_if_null = QCheckBox("潮位补全")
        self.del_data = QPushButton("删除当前潮位数据", clicked=self.del_all_data)
        self.tide_d_Group = QGroupBox("其他")
        self.tide_d2_Group = QGroupBox('文档输出')

        self.errorMessageDialog = QErrorMessage(self)
        self.errorLabel = QLabel()
        self.errorButton = QPushButton("确定")

        tide_l_Group_Layout = QVBoxLayout()
        tide_l_Group_Layout.addWidget(self.tide_file_button)
        tide_l_Group_Layout.addWidget(self.auto_Preprocess)
        tide_l_Group_Layout.addWidget(self.if_one_sheet)
        tide_l_Group_Layout.addWidget(self.tide_d_Group)
        tide_l_Group_Layout.addWidget(self.tide_d2_Group)

        self.tide_l_Group.setLayout(tide_l_Group_Layout)

        tide_d_Group_Layout = QVBoxLayout()

        tide_d_Group_Layout.addWidget(self.check_multifile)
        tide_d_Group_Layout.addWidget(self.del_data)

        self.show_statics = QPushButton('显示各站点统计结果')
        self.gen_Tieddocument_button = QPushButton('生成潮位报告', clicked=self.genTideDocument2File)

        tide_d2_Group_Layout = QVBoxLayout()
        tide_d2_Group_Layout.addWidget(self.show_statics)
        tide_d2_Group_Layout.addWidget(self.gen_Tieddocument_button)

        self.tide_d_Group.setLayout(tide_d_Group_Layout)
        self.tide_d2_Group.setLayout(tide_d2_Group_Layout)

        self.sites_Tab = QTabWidget()

        # 调试site_tab 用
        # self.sites_Tab.addTab(Tide_Site_Tab(), '未载入潮位数据')
        # self.welcome_init_tab()

        mainLayout = QHBoxLayout()
        mainLayout.addWidget(self.tide_l_Group)
        mainLayout.addWidget(self.sites_Tab)
        self.setLayout(mainLayout)

        self.sites_control = {}

    def welcome_init_tab(self):
        self.init_tab = Welcome_Tab("请载入潮位数据文件")
        self.sites_Tab.addTab(self.init_tab, '未载入潮位数据')
        self.init_tab.load_file_button.clicked.connect(self.open_file)

    def open_file(self):
        options = QFileDialog.Options()
        if self.check_multifile.isChecked():
            fileName, _ = QFileDialog.getOpenFileNames(self, "选取文件", "",
                                                       "97-03表格文档 (*.xls);;Excel文件(*.xlsx)", options=options)
        else:
            fileName, filetype = QFileDialog.getOpenFileName(self, "选取文件", "", "97-03表格文档 (*.xls);;Excel文件(*.xlsx)",
                                                             options=options)
            fileName = [fileName]
        # 测试用
        # fileName = ["E:\★★★★★CODE★★★★★\Git\codes\ocean_hydrology_data_process\测试潮位.xls"]

        for f in fileName:
            try:
                t.update({f: Process_Tide(f, only_first_sheet=self.if_one_sheet.isChecked())})

            except:
                self.errorMessage(EOFError)
            if self.auto_Preprocess.isChecked():
                for i in t.values():
                    for sitename in i.sites:
                        tide_thread = Thread(target=i.preprocess, kwargs={'s': sitename})
                        second = 0
                        tide_thread.start()
                        while tide_thread.is_alive():
                            sig.msg_to_show.emit(
                                "正在进行数据预处理，数据来自文件" + fileName[0] + ',处理用时' + str(round(second, 1)) + '秒')
                            time.sleep(0.1)
                            second += 0.1
                        sig.msg_to_show.emit(fileName[0] + '文件处理完毕')
                sig.msg_to_show.emit("文件全部处理完毕，请继续操作")

        for filename, process_tide in t.items():
            # i为文件名
            # j为process calss
            for site in process_tide.sites:
                self.sites_control.update({site: Tide_Site_Tab(site, filename)})
                self.sites_Tab.addTab(self.sites_control[site], site)
                self.sites_Tab.setTabToolTip(self.sites_Tab.count() - 1, filename)

                sig.load_tide_file_done.emit(filename, site)  # 信号绑定
                self.sites_control[site].signal_connect()
                self.sites_control[site].Save_Sheet_Button.clicked.connect(
                    self.sites_control[site].save_one_site_tide_clicked)

        if t and self.sites_Tab.tabText(0) == '未载入潮位数据':
            self.sites_Tab.removeTab(0)

    def del_all_data(self):
        MESSAGE = "将删除所有已载入的潮位数据，是否确定？"
        reply = QMessageBox.question(self, "删除潮位数据",
                                     MESSAGE,
                                     QMessageBox.Yes | QMessageBox.No)
        if reply == QMessageBox.Yes:
            t.clear()
            self.clear_tab()
            self.welcome_init_tab()
        else:
            pass
        sig.msg_to_show.emit("所载入数据已全部删除")

    def clear_tab(self):
        while self.sites_Tab.count() != 0:
            self.sites_Tab.removeTab(0)

    def errorMessage(self, Wrong_msg):
        self.errorMessageDialog.showMessage(Wrong_msg)
        self.errorLabel.setText("确定")

    def genTideDocument2File(self):
        self.tideDocumentName, _ = QFileDialog.getSaveFileName(self, "保存潮位报告文件", "", "Word 文档(*.docx)",
                                                               options=QFileDialog.Options())
        if self.tideDocumentName:
            self.genTideDocument(self.tideDocumentName)

    def genTideDocument(self, filename):
        document = Document()
        document.styles['Normal'].font.name = u'宋体'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        document.styles['Heading 2'].font.name = u'宋体'
        document.styles['Heading 2']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        h = document.add_heading('潮位分析')
        h.style.font.size = Pt(20)
        h.style.font.name = u'宋体'
        h.style.font.color.rgb = RGBColor(0, 0, 0)
        h.style.font.bold = False
        p = document.add_paragraph('')
        tideStaticsContent = []
        for i, j in t.items():
            for site in j.sites:
                tideStaticsContent.append(j.outtxt[site] + '\n')
        for i in tideStaticsContent:
            r = p.add_run(i)
            r.style.font.size = Pt(10)
            r.style.font.name = u'宋体'
            r.style.font.bold = False

        document.save(filename)




class Site_Tide_signal_event(QObject):
    tide_preprocess = pyqtSignal(str, str, float)
    show_tide_HTML_signal = pyqtSignal(str, str, str)
    show_tide_TXT_signal = pyqtSignal(str, str, str)
    show_tide_statics_signal = pyqtSignal(str, str, str, str)
    save_one_site_signal = pyqtSignal(str, str)
    show_tide_harmonic_result_signal = pyqtSignal(str, str, str)
    show_diff_signal = pyqtSignal(matplotlib.figure.Figure)


class Site_Current_signal_event(QObject):
    show_current_proper_TXT_signal = pyqtSignal(str, str, str)
    show_current_arrow_signal = pyqtSignal(matplotlib.figure.Figure)


class All_Current_signal_event(QObject):
    show_all_out_signal = pyqtSignal(str, str, str)
    show_max_out_signal = pyqtSignal(str, str, str)
    show_mean_out_signal = pyqtSignal(str, str, str)
    import_report_signal = pyqtSignal(Read_Report)


class All_Signal_event(QObject):
    load_tide_file_done = pyqtSignal(str, str)
    load_current_file_done = pyqtSignal(str, str)
    msg_to_show = pyqtSignal(str)


class Tide_Site_Tab(QWidget):
    def __init__(self, sitename, filename, parent=None):
        super(Tide_Site_Tab, self).__init__(parent)
        self.layout_element()
        self.sitename = sitename
        self.filename = filename
        self.signal = Site_Tide_signal_event()

    def layout_element(self):
        self.lu_Group = QGroupBox("测点位置")
        self.lu_Group.setFixedWidth(150)
        self.ld_Group = QGroupBox("起止时间")
        self.mid_Group_out = QGroupBox("潮位查看")
        self.mid_Group_in = QGroupBox("调整去噪度")
        self.r_Group = QGroupBox("结果输出")
        self.r_in_Group_u = QGroupBox("调和分析")
        self.r_in_Group_d = QGroupBox("查看统计结果")
        # self.add_Group1 = QGroupBox("施测信息")
        self.add_Group2 = QGroupBox("高程基准")
        self.add_Group = QGroupBox()

        """add_Group1 = QFormLayout()
        self.Survey_man = QLineEdit()
        self.Survey_Implement = QComboBox()
        self.Survey_Implement.addItem("Tide Master")
        self.Survey_Implement.addItem("其他设备（待补充）")
        self.Survey_Implement.addItem("自定义")
        self.Survey_Implement_Series_num = QLineEdit()
        add_Group1.addRow(QLabel("测量人员"),self.Survey_man)
        add_Group1.addRow(QLabel("采用仪器"), self.Survey_Implement)
        add_Group1.addRow(QLabel("仪器编号"),self.Survey_Implement_Series_num)
        self.add_Group1.setLayout(add_Group1)"""

        add_Group2 = QFormLayout()
        self.Altitudinal_System = Altitudinal_Select()
        self.Altitudinal_Add = QDoubleSpinBox()
        self.Altitudinal_Add.setRange(-100, 100)
        self.Altitudinal_Add.setToolTip("单位与原数据一致，计算方式为加")
        self.confirm_change_altitudinal = QPushButton('改变数据高程', clicked=self.change_altitude)
        self.confirm_change_altitudinal.setToolTip("单位与原数据一致，计算方式为加")
        add_Group2.addRow(QLabel("数据基准面"), self.Altitudinal_System.combobox)
        add_Group2.addRow(QLabel("改变基准面"), self.Altitudinal_Add)
        add_Group2.addRow(self.confirm_change_altitudinal)
        self.add_Group2.setLayout(add_Group2)

        add_Group = QHBoxLayout()
        # add_Group.addWidget(self.add_Group1)
        add_Group.addWidget(self.add_Group2)
        self.add_Group.setLayout(add_Group)

        lu_GroupLayout = QFormLayout()
        self.onlyDouble = QDoubleValidator()
        self.Longitude = QLineEdit()
        self.Latitude = QLineEdit()
        self.Longitude.setValidator(self.onlyDouble)
        self.Latitude.setValidator(self.onlyDouble)

        lu_GroupLayout.addRow(QLabel("经度"))
        lu_GroupLayout.addRow(self.Longitude)
        lu_GroupLayout.addRow(QLabel("纬度"))
        lu_GroupLayout.addRow(self.Latitude)
        self.lu_Group.setLayout(lu_GroupLayout)

        ld_GroupLayout = QFormLayout()
        self.start_time = QDateTimeEdit()
        self.end_time = QDateTimeEdit()
        ld_GroupLayout.addRow(QLabel("开始时间"))
        ld_GroupLayout.addRow(self.start_time)
        ld_GroupLayout.addRow(QLabel("结束时间"))
        ld_GroupLayout.addRow(self.end_time)
        self.ld_Group.setLayout(ld_GroupLayout)

        mid_Group_out = QVBoxLayout()
        self.Show_tide_data_button1 = QPushButton("&潮位查看\n（WEB模式）", clicked=self.show_tide_HTML_click)

        self.Show_tide_data_button2 = QPushButton("&潮位查看\n（普通模式）", clicked=self.show_tide_TXT_click)

        self.Show_tide_data_button3 = QPushButton("&潮位查看\n（图片）", clicked=self.show_diff_pic)

        self.Process_method_select = QComboBox()
        self.Process_method_select.addItem("傅立叶分析")
        self.Process_method_select.addItem("数值平滑")
        self.confirmButton = QPushButton("数据处理", clicked=self.preprocess_click)
        self.Process_threshold = QSpinBox()
        self.Process_threshold.setRange(1, 100)
        self.Process_threshold.setValue(30)
        mid_Group_out.addWidget(self.Show_tide_data_button1)
        mid_Group_out.addWidget(self.Show_tide_data_button2)
        mid_Group_out.addWidget(self.Show_tide_data_button3)
        mid_Group_out.addSpacing(20)
        # mid_Group_out.addWidget(self.mid_Group_in,stretch=4)
        mid_Group_in = QFormLayout()
        mid_Group_in.addRow(self.Process_method_select)
        mid_Group_in.addRow("阀值", self.Process_threshold)
        mid_Group_in.addRow(self.confirmButton)
        self.mid_Group_out.setLayout(mid_Group_out)
        self.mid_Group_in.setLayout(mid_Group_in)

        r_Group = QFormLayout()
        self.Save_Sheet_Button = QPushButton("保存潮位报表")
        self.Save_Mid_Data_Button = QPushButton("保存中间结果", clicked=self.save_one_site_mid_data)
        self.Astronomical_Tide_Analysis = QPushButton("天文潮计算", clicked=self.Astronomical_analysis)
        self.analysis_by_init_data = QCheckBox("使用原始数据计算")
        self.Theoretical_Mininum_tidal_leval = QPushButton("计算对应理论最低潮面")
        self.Theoretical_Mininum_tidal_leval.setEnabled(False)
        self.document_out_button = QPushButton("输出文字统计", clicked=self.show_statics_click)
        self.show_statics_button = QPushButton("查看统计结果")
        r_Group.addRow(self.Save_Sheet_Button, self.Save_Mid_Data_Button)
        r_Group.addRow(self.r_in_Group_u)
        r_in_Group_u = QFormLayout()
        r_in_Group_u.addRow(self.Astronomical_Tide_Analysis)
        r_in_Group_u.addRow(self.analysis_by_init_data)
        # r_in_Group_u.addRow(self.Theoretical_Mininum_tidal_leval)

        self.r_in_Group_u.setLayout(r_in_Group_u)

        r_Group.addRow(self.r_in_Group_d)
        r_in_Group_d = QFormLayout()
        # r_in_Group_d.addRow(self.show_statics_button)
        r_in_Group_d.addRow(self.document_out_button)
        self.r_in_Group_d.setLayout(r_in_Group_d)
        self.r_Group.setLayout(r_Group)

        mainLayout = QHBoxLayout()
        mainLayout.addWidget(self.lu_Group)
        mainLayout.addWidget(self.ld_Group)
        mainLayout.addWidget(self.add_Group)
        mainLayout.addWidget(self.mid_Group_out)
        mainLayout.addWidget(self.mid_Group_in)
        mainLayout.addWidget(self.r_Group)
        self.setLayout(mainLayout)

    def preprocess_click(self):
        self.signal.tide_preprocess.emit(self.filename, self.sitename, self.Process_threshold.value())

    def show_tide_HTML_click(self):
        self.generate_tide_HTML()
        self.signal.show_tide_HTML_signal.emit(self.filename, self.sitename,
                                               '处理阀值为：' + str(self.Process_threshold.value()))

    def show_diff_pic(self):
        # plot_thread = Thread(target=t[self.filename].plot_tide_compare,kwargs={'site':self.sitename})
        self.fig = t[self.filename].plot_tide_compare(self.sitename)
        self.signal.show_diff_signal.emit(self.fig)
        '''plot_thread.start()
        second = 0
        while plot_thread.is_alive():
            sig.msg_to_show.emit('正在生成交互图形,处理用时'+str(round(second,1))+'秒')
            time.sleep(0.1)
            second += 0.1'''
        sig.msg_to_show.emit(self.sitename+'站图形生成完毕')
        #self.signal.show_diff_signal.emit(t[self.filename].compare_fig)

    def show_tide_TXT_click(self):
        self.signal.show_tide_TXT_signal.emit(self.filename, self.sitename,
                                              '处理阀值为：' + str(self.Process_threshold.value()))

    def signal_connect(self):
        self.signal.save_one_site_signal.connect(self.save_one_site_tide)

    def generate_tide_HTML(self):
        html_process_thread = Thread(target=t[self.filename].display, kwargs={'site': self.sitename})
        html_process_thread.start()
        second = 0
        while html_process_thread.is_alive():
            sig.msg_to_show.emit('正在生成潮位对应HTML文件，处理用时' + str(round(second, 1)) + '秒')
            time.sleep(0.1)
            second += 0.1
        sig.msg_to_show.emit(self.filename + '文件中的' + self.sitename + '站HTML潮位显示处理完毕')

    def show_statics_click(self):
        self.signal.show_tide_statics_signal.emit('TXT', self.sitename + '潮位特征值',
                                                  t[self.filename].outtxt[self.sitename].replace(',', '\n\t\t').replace(
                                                      '的', '的统计结果如下:\n\t\t'), '数据源文件为' + self.filename)

    def save_one_site_tide_clicked(self):
        self.signal.save_one_site_signal.emit(self.filename, self.sitename)

    def save_something_to_excel(self, fun, args, file_type_in_chinese):

        options = QFileDialog.Options()
        savefileName, _ = QFileDialog.getSaveFileName(self, "保存" + file_type_in_chinese, "",
                                                      "Excel文件 (*.xlsx)", options=options)
        if savefileName:
            args.update({'outfile': savefileName})
            save_tide_thread = Thread(target=fun,
                                      kwargs=args)
            second = 0
            save_tide_thread.start()
            while save_tide_thread.is_alive():
                sig.msg_to_show.emit(
                    '正在生成' + self.sitename + file_type_in_chinese + '，已用时' + str(round(second, 1)) + '秒，请稍等')
                time.sleep(0.1)
                second += 0.1
            sig.msg_to_show.emit(self.sitename + file_type_in_chinese + '文件' + savefileName + '生成结束')
            self.filename = savefileName

    def save_one_site_mid_data(self):
        self.save_something_to_excel(t[self.filename].out_put_mid_data, {'sitename': self.sitename}, '计算过程文件')

    def save_one_site_tide(self):
        self.save_something_to_excel(t[self.filename].output, {'sitename': self.sitename,
                                                               'time_start': self.start_time.dateTime().toPyDateTime()
            , 'time_end': self.end_time.dateTime().toPyDateTime()
            , 'reference_system': self.Altitudinal_System.combobox.currentText()}, '潮位报表')

    def change_altitude(self):
        t[self.filename].change_altitude(self.sitename, diff=self.Altitudinal_Add.value())

        sig.msg_to_show.emit('高程处理结束')

    # def tideanalysis(self,sitename),:

    def Astronomical_analysis(self):
        harmonic_thread = Thread(target=t[self.filename].harmonic_analysis,
                                 kwargs={'site': self.sitename, 'if_init': self.analysis_by_init_data.isChecked()})
        harmonic_thread.start()
        second = 0
        while harmonic_thread.is_alive():
            sig.msg_to_show.emit('正在进行调和分析，处理用时' + str(second) + '秒')
            time.sleep(0.2)
            second += 0.2
        sig.msg_to_show.emit(self.filename + '文件中的' + self.sitename + '站调和分析处理完毕')
        if self.analysis_by_init_data.isChecked():
            s = '原始数据'
        else:
            s = '去噪后的数据'
        self.signal.show_tide_harmonic_result_signal.emit(self.filename, self.sitename, s)
        self.Theoretical_Mininum_tidal_leval.setEnabled(True)


class Altitudinal_Select(QDialog):
    def __init__(self, parent=None):
        super(Altitudinal_Select, self).__init__(parent)
        self.combobox = QComboBox()
        self.combobox.addItem("当地理论最低潮面")
        self.combobox.addItem("平均海平面")
        self.combobox.addItem("85高程")
        self.combobox.addItem("自定义")
        self.combobox.activated.connect(self.changed)

    def changed(self, index):
        if index == 0:
            self.select = "当地理论最低潮面"
        elif index == 1:
            self.select = "平均海平面"
        elif index == 2:
            self.select = "85高程"
        elif index == 3:
            self.setText()

    def setText(self):
        # text, ok = QInputDialog.getText(self, "自定义高程系统", "自定义名称:", QLineEdit.Normal)
        text, ok = QInputDialog.getText(self, "自定义高程系统",
                                        "自定义名称:", QLineEdit.Normal, self.combobox.itemText(3))
        if ok and text != '':
            self.select = text
            self.combobox.addItem(text)
            self.combobox.removeItem(3)


class Welcome_Tab(QWidget):
    def __init__(self, message):
        super(Welcome_Tab, self).__init__(parent=None)
        self.load_file_button = QPushButton(message)
        self.load_file_button.setStyleSheet("""
               QPushButton
               {
                   background-color: rgb(255, 255,255);
                   border:10px solid rgb(0, 170, 255);
               }
               """)
        main_Layout = QFormLayout()
        main_Layout.addRow(self.load_file_button)
        self.setLayout(main_Layout)


class Current_Tab(QWidget):
    def __init__(self, parent=None):
        super(Current_Tab, self).__init__(parent)
        global c
        c = {}
        self.layout_element()
        self.errorMessageDialog = QErrorMessage(self)
        self.options = QFileDialog.Options()
        self.beding_current = []
        self.Point_control = {}
        self.list_of_all_current_out = []
        self.all_current_out = None
        # self.welcome_init_tab()
        self.sig = All_Current_signal_event()
        self.filename = None

    def bed_group_Layout(self):
        r_group_layout = QFormLayout()
        self.pre_bed = QPushButton('导入文件', clicked=self.open_bed_file)
        self.pre_bed.setToolTip('载入潮流文件应包括时间、各层的流速数据（两种格式均可），其中流速单位为m/s')
        """self.check_multifile_bed = QCheckBox("多个文件")
        self.check_multifile_bed.setToolTip('按照所设置的高度同时转化多个文件')"""
        self.if_vd = QCheckBox("V-D格式")
        self.if_vd.setChecked(False)
        self.if_vd.setToolTip("流速为V-D格式时选择此项，为东北分向流速时取消选择")

        self.if_reverse = QCheckBox('沉底式')
        self.if_reverse.setChecked(False)
        self.if_reverse.setToolTip('采用沉底式ADCP时选择此项，层数由底层到表层')

        self.bed_count = QComboBox()
        self.bed_count.addItem('六层法')
        self.bed_count.addItem('三层法（尚未开发）')
        self.first_bin = QDoubleSpinBox()
        self.first_bin.setValue(0.7)
        self.first_bin.setToolTip("输入仪器探头在水下的深度加上盲区深度，单位为米")
        self.bin_height = QDoubleSpinBox()
        self.bin_height.setValue(1.0)
        self.bin_height.setToolTip('设置单层高度，单位为米')

        self.top_ratio = QDoubleSpinBox()
        self.top_ratio.setValue(1.05)
        self.top_ratio.setRange(1, 1.2)
        self.top_ratio.setSingleStep(0.01)

        self.button_ratio = QDoubleSpinBox()
        self.button_ratio.setValue(0.95)
        self.button_ratio.setRange(0.9, 1)
        self.button_ratio.setSingleStep(0.01)

        self.mag_declination = QDoubleSpinBox()
        self.mag_declination.setRange(-180, 180)
        self.mag_declination.setValue(1.91)#吉布提
        #self.mag_declination.setValue(-5.80)
        #self.mag_declination.setValue(-5.81)#东海舟山

        self.mag_declination.setSingleStep(0.01)
        self.mag_declination.setToolTip('磁偏角向东为正向西为负，默认值1.91为吉布提对应磁偏角')


        self.if_save_VD = QCheckBox('保存为格式')
        self.if_save_VD.setChecked(True)
        self.save_bedded_file = QPushButton('分层并保存结果', clicked=self.save_bedded_file_click)

        r_group_layout.addRow(self.pre_bed)
        r_group_layout.addRow(self.if_vd, self.if_reverse)
        # r_group_layout.addRow(QLabel("总层数选择"))
        # r_group_layout.addRow(self.bed_count)
        r_group_layout.addRow("入水深度加盲区高度", self.first_bin)
        r_group_layout.addRow("每层高度", self.bin_height)
        r_group_layout.addRow("表层流速比例", self.top_ratio)
        r_group_layout.addRow("底层流速比例", self.button_ratio)
        r_group_layout.addRow("磁偏角", self.mag_declination)
        r_group_layout.addRow(self.save_bedded_file)
        r_group_layout.addRow(self.VD2EN)
        r_group_layout.addRow(self.EN2VD)

        self.bed_group_layout = r_group_layout

    def layout_element(self):

        self.current_l_Group = QGroupBox('总览')
        self.current_l_Group.setMaximumWidth(230)

        self.current_m_Group = QGroupBox('文件预处理')
        self.current_m_Group.setMaximumWidth(200)
        self.bed_Group = QGroupBox("文件分层")

        self.del_data = QPushButton("删除当前潮流数据", clicked=self.del_all_data)
        #self.save_imported_file_button = QPushButton("保存载入数据",clicked = self.save_imported_file)
        #self.import_files_button = QPushButton("保存载入数据", clicked=self.import_files)

        self.VD2EN = QPushButton('合成流速转为分向流速', clicked=self.vd_to_en)
        self.EN2VD = QPushButton('分向流速转为合成流速', clicked=self.ne_to_vd)

        self.Point_count = QSpinBox()
        self.Point_count.setRange(1, 100)
        self.Point_count.setValue(1)
        self.Point_count.setMaximumWidth(40)

        self.all_Tide_type = QComboBox()
        self.all_Tide_type.addItem("大、中、小三潮")
        self.all_Tide_type.addItem("大、小两潮")

        self.ceng_select = QComboBox()
        self.ceng_select.addItem('六点法')
        self.ceng_select.addItem('三点法')
        self.save_Point_and_Tidetype_count = QPushButton('确认测点、潮型', clicked=self.confirm_point)
        self.import_from_report = QPushButton('从潮流报表导入',clicked = self.import_from_report_click)
        self.change_Point_name = QLineEdit()
        self.change_Point_name.setMaximumHeight(15)



        self.show_all_out_Button = QPushButton("查看所有测流结果", clicked=self.show_all_out)
        self.save_all_out_Button = QPushButton('保存', clicked=self.save_all_out)
        self.save_max_out_Button = QPushButton("汇总流速最大值", clicked=self.save_max_out)
        self.save_mean_out_Button = QPushButton("汇总流速平均值", clicked=self.save_mean_out)
        self.draw_current_cad_Button = QPushButton("绘制流速矢量图",clicked= self.draw_current_cad)
        self.gen_all_duration_Button = QPushButton("查看涨落潮历时",clicked = self.gen_all_duration_click)
        self.draw_cengshu = QComboBox()
        self.draw_cengshu.addItems(["垂线平均","表层","0.2H","0.4H","0.6H","0.8H","底层"])
        self.cad_parameter = QSpinBox()
        self.cad_parameter.setValue(10000)




        l_group_layout = QFormLayout()
        l_group_layout.addRow("选择潮流测点数量", self.Point_count)
        l_group_layout.addRow("选择测流潮型数", self.all_Tide_type)
        l_group_layout.addRow("选择垂线分层数", self.ceng_select)
        l_group_layout.addRow(self.save_Point_and_Tidetype_count,self.import_from_report)
        l_group_layout.addRow(QLabel("编辑测点名称"))
        l_group_layout.addRow(self.change_Point_name)
        l_group_layout.addRow(self.save_mean_out_Button,self.save_max_out_Button)
        l_group_layout.addRow(self.show_all_out_Button, self.save_all_out_Button)
        l_group_layout.addRow(self.draw_current_cad_Button,self.draw_cengshu)
        l_group_layout.addRow("选择绘图比例",self.cad_parameter)
        l_group_layout.addRow(self.del_data)

        self.current_l_Group.setLayout(l_group_layout)

        m_group_layout = QVBoxLayout()

        self.bed_group_Layout()
        self.bed_Group.setLayout(self.bed_group_layout)

        m_group_layout.addWidget(self.bed_Group)

        self.current_m_Group.setLayout(m_group_layout)

        self.sites_tab = QTabWidget()

        mainLayout = QHBoxLayout()
        mainLayout.addWidget(self.current_m_Group, stretch=20)
        mainLayout.addWidget(self.current_l_Group, stretch=20)
        mainLayout.addWidget(self.sites_tab, stretch=50)
        self.setLayout(mainLayout)

    def gen_all_current_out(self):

        try:
            if len(c) != self.Point_count.value() * len(self.Tide_types) or not c:
                QMessageBox.information(self, '通知',
                                        "请载入所有点位对应数据后再点击")


            elif not self.list_of_all_current_out:
                for _, j in c.items():
                    self.list_of_all_current_out.append(j.output_all())
                self.all_current_out = pandas.concat(self.list_of_all_current_out)

        except:
            QMessageBox.information(self, '通知',
                                    "请载入所有点位对应数据后再点击")

    def gen_all_duration_out(self):
        self.gen_all_current_out()
        self.dic_duration = {}
        for i,j in c.items():
            self.dic_duration.update({i:j.out_times()})

    def ave_out(self):
        def str_type(point_type_re):
            s = ''
            s += '测点' + point_type_re[0] + '在' + point_type_re[1] + point_type_re[2] + '时'

        def out_ave_txt_c(row):
            return str_from_df(row['层数']) + '的平均流速为' + num_from_df(row['平均流速']) + 'cm/s,平均流向为' + num_from_df(
                row['平均流向']) + '°,'

        try:
            g3 = self.all_current_out.groupby(['Point', '潮型', '涨落潮'])
            l_cengshu = ['表层', '0.2H', '0.4H', '0.6H', '0.8H', '底层', '垂线平均']
            s = ''
            for i, j in g3:
                s += str_type(i)
                for ceng in l_cengshu:
                    s += out_ave_txt_c(j[j['层数'] == ceng])
            return s
        except:
            self.gen_all_current_out()
            self.ave_out()
        self.gen_all_current_out()

    def gen_all_duration_click(self):
        self.sig.show_all_out_signal.emit('TXT','涨落潮历时汇总',pandas.DataFrame(self.dic_duration).to_string())

    def save_csv(self, df):
        options = QFileDialog.Options()
        filename, _ = QFileDialog.getSaveFileName(self, "保存所有点位特征值文件", '', "Excel文件(*.xlsx);;逗号分割文档(utf-8)(*.csv)",
                                                  options=options)
        if filename:
            self.filename = filename
            if "csv" in filename:
                df.to_csv(filename)
            else:
                df.to_excel(filename)

    def show_all_out(self):
        self.gen_all_current_out()
        self.all_current_out = self.all_current_out[
            ['Point', '潮型', '层数', '涨落潮', '平均流速', '平均流向', '最大流速', '最大流速对应方向', '出现时间', '来源文件']]

        self.sig.show_all_out_signal.emit('TXT', '潮流计算结果汇总', self.all_current_out.to_string())
        ######

    def save_all_out(self):
        self.gen_all_current_out()
        self.save_csv(self.all_current_out)

    def save_max_out(self):
        self.gen_all_current_out()
        self.all_current_out['最大流速对应方向'] = pandas.to_numeric(self.all_current_out['最大流速对应方向'])
        self.all_current_out['最大流速'] = pandas.to_numeric(self.all_current_out['最大流速'])
        self.max_out = pandas.pivot_table(self.all_current_out, values=['最大流速', '最大流速对应方向'],
                                          index=['Point', '潮型', '涨落潮'], columns=['层数']).swaplevel(0, 1,
                                                                                                  axis=1).reindex(
            ['垂线平均', '表层', '0.2H', '0.4H', '0.6H', '0.8H', '底层'], axis=1, level=0).reindex(['最大流速', '最大流速对应方向'],
                                                                                                axis=1, level=1).reindex(['大潮','中潮','小潮'],axis = 0,level=1)
        self.max_out = self.max_out  # 改一下层数顺序，下同
        self.save_csv(self.max_out)
        sig.msg_to_show.emit('流速最大值汇总结束')

    def save_mean_out(self):
        self.gen_all_current_out()
        self.all_current_out['平均流速'] = pandas.to_numeric(self.all_current_out['平均流速'])
        self.all_current_out['平均流向'] = pandas.to_numeric(self.all_current_out['平均流向'])
        self.mean_out = pandas.pivot_table(self.all_current_out, values=['平均流速', '平均流向'], index=['Point', '潮型', '涨落潮'],
                                           columns=['层数']).swaplevel(0, 1, axis=1).reindex(
            ['垂线平均', '表层', '0.2H', '0.4H', '0.6H', '0.8H', '底层'], axis=1, level=0).reindex(['平均流速', '平均流向'],
                                                                                                axis=1, level=1).reindex(['大潮','中潮','小潮'],axis = 0,level=1)

        self.mean_out = self.mean_out  # 改一下层数顺序，下同
        self.save_csv(self.mean_out)
        sig.msg_to_show.emit('流速平均值汇总结束')

    def welcome_init_tab(self):
        self.init_tab = Welcome_Tab("请载入单点潮流数据文件")
        # 测试单个测点tab
        self.sites_tab.addTab(self.init_tab, '未载入潮流数据')
        # self.sites_tab.addTab(Point_Tab("POINT1", tide_types=['大潮', '中潮', '小潮'], cengshu=6), 'POINT1')
        self.init_tab.load_file_button.clicked.connect(self.open_current_file)

    def confirm_point(self,format_report = False):
        if self.all_Tide_type.currentIndex() == 0:
            self.Tide_types = ['大潮', '中潮', '小潮']
        else:
            self.Tide_types = ['大潮', '小潮']
        if self.ceng_select.currentIndex() == 0:
            self.ceng_count = 6
        else:
            self.ceng_count = 3
        try:
            print(str(self.Points))
        except:
            point_names = self.change_Point_name.text()
            if point_names:
                point_names = point_names.split(sep=',')
                if len(point_names) != self.Point_count.value():
                    QMessageBox.information(self, '通知',
                                            "请将站点名称用“,”分隔，并与测点数量相对应")
                    return None

                self.Points = point_names
            else:
                self.Points = ['P' + str(x) for x in range(1, self.Point_count.value() + 1)]

        if self.sites_tab.tabText(0) == '未载入潮流数据':
            self.sites_tab.removeTab(0)

        for i in self.Points:
            self.Point_control_creat(i)

    def process_excel_data(self,c_data):
        QApplication.processEvents()
        threads = []
        for _,i in c_data.data.items():
            for _,ii in i.items():
                one_thread = Process_excel_data_thread(ii)
                threads.append(one_thread)
        for one in threads:
            one.start()
        #report_data_process = QThread()
        second = 0

        print(threads[-1].isFinished())

        while not threads[-1].isFinished():
            time.sleep(1)
            print(second)
            second +=1
            pass
        sig.msg_to_show.emit('各测站预处理数据结束，请稍等')
        print('ok~! '*10)
        return 0

    def import_from_report_click(self):

        self.current_report_file,_ = QFileDialog.getOpenFileName(self,"选取潮流报表文件","","Excel文件(*.xlsx)")
        if not self.current_report_file:
            return 0
        c_excel = Read_Report(self.current_report_file)
        #self.Points = list(c2.points)

        self.Point_count.setValue(len(c_excel.points))
        if len(c_excel.tide_type) == 3:
            self.all_Tide_type.setCurrentIndex(0)
        else:
            self.all_Tide_type.setCurrentIndex(1)
        self.Points = list(c_excel.points)
        self.confirm_point()

        #设置各点的流向
        Excel_angles = {}
        for P_name in c_excel.points:
            ang = self.set_Angles_of_excel_data(P_name)
            Excel_angles.update({P_name:ang})
            c_excel.setPoint_ang(P_name,ang = ang)

        sig.msg_to_show.emit('正在处理潮流数据，请稍候')
        #读取各自ang结束
        self.process_excel_data(c_excel) #流速数据处理

        for Point_Name,Contral in self.Point_control.items():
            for tide in c_excel.tide_type:
                Contral.One_tide_file_open(is_format_reprot=True,reportFile=self.current_report_file,tideType = tide,angle=Excel_angles[Point_Name],data = c_excel.data[Point_Name][tide])

        sig.msg_to_show.emit('潮流报表数据处理结束')
        print('ok')

    def set_Angles_of_excel_data(self,Point_Name):
            d_angle, ok = QInputDialog.getDouble(self,"请输入"+Point_Name+"主流向",'落潮流向为正值，涨潮流向为负值',0,-360,360,2)
            if ok:
                sig.msg_to_show.emit('成功将测点'+Point_Name+'主流向设置成为'+ str(d_angle))
            return d_angle

    def Q_Table(self):
        self.tableWidget = QTableWidget()
        self.tableWidget.setRowCount(self.Point_count.value())
        self.tableWidget.setColumnCount(2)
        for i in range(self.Point_count.value()):
            self.tableWidget.setItem(i,0,QTableWidgetItem('a'))
        self.tableWidget.show()
    def Point_control_creat(self, Point):
        one_point_tab = Point_Tab(Point, tide_types=self.Tide_types, cengshu=self.ceng_count)
        self.Point_control.update({Point: one_point_tab})
        self.sites_tab.addTab(one_point_tab, Point)

    def open_bed_file(self):
        try :
            del self.beding_current
        except:
            pass
        self.pre_bed_fileName, _ = QFileDialog.getOpenFileName(self, "选取待分层文件", "", "Excel文件(*.xlsx);;逗号分割文档(utf-8)(*.csv)",
                                                               options=self.options)
        if self.pre_bed_fileName:
            self.beding_current = Current_pre_process(filename=self.pre_bed_fileName, is_VD=self.if_vd.isChecked())
            sig.msg_to_show.emit('文件载入成功')

    def save_bedded_file_click(self):
        fileName, _ = QFileDialog.getSaveFileName(self, '保存分层文件路径', self.pre_bed_fileName, "Excel文件(*.xlsx);;逗号分割文档(utf-8)(*.csv)",
                                                  options=self.options)
        if fileName:
            self.beding_current.all_fenceng(bin=self.bin_height.value(), first_bin=self.first_bin.value(),
                                            top_ratio=self.top_ratio.value(), button_ratio=self.button_ratio.value(),
                                            reverse=self.if_reverse.isChecked())
            self.beding_current.save_result(outfile=fileName, V_D=self.if_save_VD,
                                            mag_declination=self.mag_declination.value())
            QMessageBox.information(self, '通知',
                                    "文件转化结束")
            sig.msg_to_show.emit("文件转化结束")

    def open_current_file(self):
        fileName, _ = QFileDialog.getOpenFileName(self, "选取文件", '', "Excel文件(*.xlsx);;逗号分割文档(utf-8)(*.csv)", options=self.options)
        if fileName:
            self.filename = fileName
            return fileName

    def clear_tab(self):
        while self.sites_tab.count() != 0:
            self.sites_tab.removeTab(0)

    def del_all_data(self):
        MESSAGE = "将删除所有已载入的潮流数据，是否确定？"
        reply = QMessageBox.question(self, "删除潮位数据",
                                     MESSAGE,
                                     QMessageBox.Yes | QMessageBox.No)
        if reply == QMessageBox.Yes:
            c.clear()
            self.clear_tab()
            # self.welcome_init_tab()


        else:
            pass
        sig.msg_to_show.emit("所载入潮流数据已全部删除")

    def vd_to_en(self):
        options = QFileDialog.Options()
        vd_fileName, _ = QFileDialog.getOpenFileName(self, "选取流速流向格式潮流文件", "", "逗号分割文档(utf-8) (*.csv);;Excel文件(*.xlsx)",
                                                     options=options)
        ne_fileName, _ = QFileDialog.getSaveFileName(self, "保存为分向流速格式潮流文件", "", "Excel文件(*.xlsx);;逗号分割文档(utf-8)(*.csv)",
                                                     options=options)
        try:
            vd = Current_pre_process(filename=vd_fileName, is_VD=True)
            vd.convert_to_EN_file(ne_fileName)
            QMessageBox.information(self, '通知',
                                    "文件转化结束")
        except:
            pass

    def ne_to_vd(self):
        options = QFileDialog.Options()
        ne_fileName, _ = QFileDialog.getOpenFileName(self, "选取分向流速格式潮流文件", "", "逗号分割文档(utf-8) (*.csv);;Excel文件(*.xlsx)",
                                                     options=options)
        vd_fileName, _ = QFileDialog.getSaveFileName(self, "保存为流速流向格式潮流文件", dir_of_file(ne_fileName), "Excel文件(*.xlsx);;逗号分割文档(utf-8)(*.csv)",
                                                     options=options)
        try:
            ne = Current_pre_process(filename=ne_fileName, is_VD=False)
            ne.convert_to_VD_file(vd_fileName)
            QMessageBox.information(self, '通知',
                                    "文件转化结束")
        except:
            pass

    def draw_current_cad(self):
        self.gen_all_current_out()

        def end_point(x, y, velocity, direction, parameter):
            v_east = velocity * np.sin(direction / 180 * np.pi)
            v_north = velocity * np.cos(direction / 180 * np.pi)
            return [x + v_east * parameter, y + v_north * parameter]

        def plot_line(x, y, vs, ds, tuceng, parameter,drawing):
            for v, d in zip(vs, ds):
                line = dxf.line((x, y), end_point(x, y, v, d, parameter))
                line['layer'] = tuceng
                drawing.add(line)
                # layer_name = dxf.layer(cengshu+self.tide_type)
                # drawing.layers.add(layer_name)
        def select_positions_file(self):        
            self.position_file,_ = QFileDialog.getOpenFileName(self,"选取位置文件","","Excel文件(*.xlsx);;逗号分割文档(utf-8)(*.csv)")
            if "csv" in self.position_file:
                self.positions = pandas.read_csv(self.position_file,usecols=['name','x','y'])
            else:
                self.positions = pandas.read_excel(self.position_file, columns=['name', 'x', 'y'])
        try:
            if(self.positions):
                pass
        except:
            self.select_positions_file()

        self.cad_file,_ = QFileDialog.getSaveFileName(self,"选取cad文件保存路径","","DXF文档(*.dxf)")
        drawing = dxf.drawing(self.cad_file)
        named_points = []
        for pos in self.positions.itertuples():
            c_of_one_point = [point_c for _,point_c in c.items() if point_c.point == pos.name]
            for every_c in c_of_one_point:
                every_c.location(x = pos.x, y = pos.y)
                data = every_c.ceng_processed[self.draw_cengshu.currentIndex()].data
                data = data[data['t'].apply(lambda t: t.minute) == 0]
                plot_line(every_c.x,every_c.y, data['v'], data['d'],tuceng= every_c.tide_type+self.draw_cengshu.currentText(), parameter=self.cad_parameter.value(),drawing=drawing)

            if every_c.point not in named_points:
                text = dxf.text(every_c.point,insert=(every_c.x,every_c.y),height=self.cad_parameter.value()*10,)
                text['layer'] = ['TEXT']
                text['color'] = 5
                drawing.add(text)
                named_points.append(every_c.point)
        drawing.save()
        QMessageBox.information(self,'通知',"流速矢量图绘制完成")

class Wind_Tab(QWidget):
    def __init__(self, parent=None):
        super(Wind_Tab, self).__init__(parent)
        self.layout_element()

    def layout_element(self):
        self.load_file = QPushButton('载入测风数据', clicked=self.preprocess_wind)
        self.load_file.setFixedHeight(50)
        self.statics = QPushButton('测风数据统计分析')
        self.statics.setFixedHeight(50)
        self.fig = QPushButton('查看各风向频率分布图')
        self.fig.setFixedHeight(50)

        self.to_document = QPushButton('生成气象报告')
        self.to_document.setFixedHeight(50)

        mainLayout = QVBoxLayout()
        mainLayout.addWidget(self.load_file)
        mainLayout.addWidget(self.statics)
        mainLayout.addWidget(self.fig)
        mainLayout.addWidget(self.to_document)

        self.setLayout(mainLayout)

    def preprocess_wind(self):
        toCompileFilenames = QFileDialog.getOpenFileNames(self, '原始气象数据', "", "DAT文件（*.dat)",
                                                          options=QFileDialog.Options())
        toSaveFilenames = QFileDialog.getSaveFileName(self, '转化后的气象数据","","逗号分割文档(utf-8)(*.csv)', options=QFileDialog.Options())
        # process_file(toCompileFilenames,toSaveFilenames)
class Wave_tab(QWidget):
    def __init__(self, parent=None):
        super(Wave_tab, self).__init__(parent)
        self.layout_element()

    def layout_element(self):
        self.load_file = QPushButton('载入波浪数据')
        self.load_file.setFixedHeight(50)
        self.statics = QPushButton('波浪数据统计分析')
        self.statics.setFixedHeight(50)
        self.fig = QPushButton('查看各波向频率分布图')
        self.fig.setFixedHeight(50)
        self.joint_distribution = QPushButton('查看各波高-周期联合分布图')
        self.joint_distribution.setFixedHeight(50)
        self.to_document = QPushButton('生成对应波浪报告')
        self.to_document.setFixedHeight(50)

        mainLayout = QVBoxLayout()
        mainLayout.addWidget(self.load_file)
        mainLayout.addWidget(self.statics)
        mainLayout.addWidget(self.joint_distribution)
        mainLayout.addWidget(self.fig)
        mainLayout.addWidget(self.to_document)

        self.setLayout(mainLayout)


class Class_Site_Tab(QWidget):
    def __init__(self, sitename, filename, parent=None):
        super(Class_Site_Tab, self).__init__(parent)
        self.layout_element()

    def layout_element(self):
        self.lu_Group = QGroupBox("测点位置")


class Point_Tab(QWidget):
    def __init__(self, Point, cengshu, tide_types=None, parent=None):
        super(Point_Tab, self).__init__(parent)
        self.point = Point
        self.tide_types = tide_types
        self.cengshu = cengshu
        self.layout_element()
        self.tides_control = {}
        self.filename = 'E:'

    def layout_element(self):
        l_Group = QGroupBox("流向设置")
        l_Group.setMaximumWidth(200)

        self.angle = QDoubleSpinBox()
        self.angle.setRange(0, 359.99)
        self.zhang_or_luo = QCheckBox('涨潮流')
        self.zhang_or_luo.setToolTip("打勾时输入为涨潮主流向，否则为落潮流向")
        self.zhang_or_luo.setChecked(False)
        self.load_Spring_file = QPushButton('载入大潮文件', clicked=self.Spring_file_open)
        self.load_Mid_file = QPushButton('载入中潮文件', clicked=self.Mid_file_open)
        self.load_Neap_file = QPushButton('载入小潮文件', clicked=self.Neap_file_open)
        self.clear_file_button = QPushButton('清除站点文件', clicked=self.clear_file)


        l_GroupLayout = QFormLayout()
        l_GroupLayout.addRow(self.angle)
        l_GroupLayout.addRow(self.zhang_or_luo)
        l_GroupLayout.addRow(self.load_Spring_file)
        if len(self.tide_types) == 3:
            l_GroupLayout.addRow(self.load_Mid_file)
        l_GroupLayout.addRow(self.load_Neap_file)


        l_GroupLayout.addRow(self.clear_file_button)

        l_Group.setLayout(l_GroupLayout)

        self.tabs = QTabWidget()  # 单潮型处理
        layout = QHBoxLayout()
        layout.addWidget(l_Group)
        layout.addWidget(self.tabs)
        self.setLayout(layout)

    def One_tide_file_open(self,is_format_reprot = False,ZhangorLuo = False,reportFile = None,tideType = None,angle = None,data = None):
        if not is_format_reprot:
            options = QFileDialog.Options()
            fileName, _ = QFileDialog.getOpenFileName(self, "选取文件", self.filename[:-8], "Excel文件(*.xlsx);;逗号分割文档(utf-8)(*.csv)", options=options)
            if fileName:
                self.filename = fileName

                def generate_tab(tide_type, **kwargs):
                    sig.msg_to_show.emit('正在载入' + tide_type + '文件,请稍候')
                    tab = Single_Tide_tab(filename=fileName, tide_type=tide_type, angle=self.angle.value(),
                                          zhang_or_luo=self.zhang_or_luo.isChecked(), Point=self.point,
                                          cengshu=self.cengshu)
                    self.tides_control.update({tide_type: tab})
                    self.tabs.addTab(tab, tide_type)
                    self.tabs.setTabToolTip('源文件为' + fileName)
                    sig.load_current_file_done.emit(self.point, tide_type)
                    sig.msg_to_show.emit(self.point + '站' + tide_type + '数据文件载入结束，请继续操作')

                return generate_tab
            else:
                x = lambda xx: print(xx)
                return x

        else:#从报表中导入,主要参数均从函数入口传递
            tab = Single_Tide_tab(filename=reportFile, tide_type=tideType, angle=angle,
                                  zhang_or_luo=ZhangorLuo, Point=self.point,
                                  cengshu=self.cengshu,is_format_reprot = is_format_reprot,data = data)
            self.tides_control.update({tideType: tab})
            self.tabs.addTab(tab, tideType)
            sig.load_current_file_done.emit(self.point,tideType)
            #self.tabs.setTabToolTip('源文件为' + reportFile)

    def clear_file(self):
        self.tides_control.clear()
        self.tabs.clear()
        x = []
        for i in c.keys():
            if self.point in i:
                x.append(i)
        for i in x:
            c.pop(i)
            sig.msg_to_show.emit(i + '已经清除')
        sig.msg_to_show.emit(self.point + '所载入文件已全部清除')

    def Spring_file_open(self):
        self.One_tide_file_open()('大潮')

    def Neap_file_open(self):

        self.One_tide_file_open()('小潮')

    def Mid_file_open(self):
        self.One_tide_file_open()("中潮")


class Single_Tide_tab(QWidget):#需根据从报表导入重构
    def __init__(self, filename, Point, cengshu, tide_type, angle, zhang_or_luo, parent=None,is_format_reprot = False,data = None):
        super(Single_Tide_tab, self).__init__(parent)
        self.Point = Point
        self.cengshu = cengshu
        self.signal = Site_Current_signal_event()

        self.tide_type = tide_type
        self.angle = angle
        self.zhang_or_luo = zhang_or_luo
        self.filename = filename

        if is_format_reprot:
            c.update({self.Point + self.tide_type: data})

        else:
            c.update({self.Point + self.tide_type: Single_Tide_Point(filename=self.filename, tide_type=tide_type,
                                                                     point=Point, angle=angle, zhang_or_luo=zhang_or_luo,
                                                                     cengshu=cengshu)})
            preprocess_thread = Thread(target=c[self.Point + self.tide_type].preprocess)
            second = 0
            preprocess_thread.start()
            while preprocess_thread.is_alive():
                sig.msg_to_show.emit('正在对' + self.tide_type + '文件进行预处理，用时' + str(round(second, 1)) + '秒')
                time.sleep(0.1)
                second += 0.1

        self.layout_element()

    def layout_element(self):
        self.output_Button = QPushButton("输出潮流特征值", clicked=self.show_current_proper_TXT_click)
        self.output_file_Button = QPushButton("输出潮流特征值文件", clicked=self.save_output)
        self.draw_stream_Button = QPushButton("查看各层逐时流速矢量图", clicked=self.show_current_arrow)

        layout = QFormLayout()
        layout.addRow(self.output_Button)
        layout.addRow(self.output_file_Button)
        layout.addRow(self.draw_stream_Button)

        self.setLayout(layout)

    def show_current_proper_TXT_click(self):
        txt = c[self.Point + self.tide_type].output_proper_txt()
        self.signal.show_current_proper_TXT_signal.emit('TXT', self.Point + self.tide_type, txt)

    def show_current_arrow(self):
        fig = c[self.Point + self.tide_type].display()
        self.signal.show_current_arrow_signal.emit(fig)

    def save_output(self):
        out = c[self.Point + self.tide_type].output_all()
        options = QFileDialog.Options()
        savefileName, _ = QFileDialog.getSaveFileName(self, "保存" + self.Point + self.tide_type + "特征值", "",
                                                      "Excel文件(*.xlsx);;逗号分割文档(utf-8)(*.csv)", options=options)
        if savefileName:
            if "csv" in savefileName:
                out.to_csv(savefileName)
            else:
                out.to_excel(savefileName)

class Process_excel_data_thread(QThread):
    def __init__(self,one_tide_point):
        QThread.__init__(self)
        self.data = one_tide_point
    def run(self):
        self.data.preprocess()

if __name__ == "__main__":

    app = QApplication(sys.argv)
    app.setStyle('Fusion')
    ex = M_window()
    ex.show()
    sys.exit(app.exec_())
