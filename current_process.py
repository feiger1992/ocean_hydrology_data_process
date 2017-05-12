#coding:utf-8
import pandas
import docx
from dxfwrite import DXFEngine as dxf
from functools import reduce
import numpy as np
import operator
import datetime
import pycnnum
import lunardate
from vectors import Vector
import xlsxwriter
import os

east = lambda v,d : v * np.sin(d/180*np.pi)
north = lambda v,d : v * np.cos(d/180*np.pi)
velocity = lambda v_e,v_n: np.sqrt(v_e**2+v_n**2)
direction = lambda v_e,v_n : 180/np.pi*np.arctan2(v_e,v_n) if(v_e>0 )else   180/np.pi*np.arctan2(v_e,v_n) +360
l_ceng=['表层','0.2H','0.4H','0.6H','0.8H','底层','垂线平均']

def plot_lines(data):#cad
    def plot_line(x, y, vs, ds, cengshu):
        for v, d in zip(vs, ds):
            line = dxf.line((x, y), end_point(x, y, v, d))
            drawing.add(line)
            line['layer'] = cengshu

    def end_point(x, y, velocity, direction):
        v_east = velocity * np.sin(direction / 180 * np.pi)
        v_north = velocity * np.cos(direction / 180 * np.pi)
        return [x + v_east, y + v_north]

    style = data['tidestyle']
    drawing = dxf.drawing('流速矢量图.dxf')
    x = data['x'].dropna().values[0]
    y = data['y'].dropna().values[0]
    plot_line(x, y, data['v_0'].values, data['d_0'], '表层')
    plot_line(x, y, data['v_2'].values, data['d_2'], '0.2层')
    plot_line(x, y, data['v_4'].values, data['d_4'], '0.4层')
    plot_line(x, y, data['v_6'].values, data['d_6'], '0.6层')
    plot_line(x, y, data['v_8'].values, data['d_8'], '0.8层')
    plot_line(x, y, data['v_10'].values, data['d_10'], '底层')
    plot_line(x, y, data['v_max'].values, data['d_v_max'], '底层')
    plot_line(x, y, data['v_max'].values, data['d_v_max'], '底层')
    plot_line(x, y, data['v'].values, data['d'], '垂线平均')
    drawing.save()
    print('ok')

def get_point_data(filename,name_of_point,data_init=None):
    print('数据来自文件'+filename)
    exl = pandas.ExcelFile(filename)
    file_data = {}
    for i,j in enumerate(exl.sheet_names):
        file_data.update({j:exl.parse(i)})
    if data_init == None:
        return {name_of_point:file_data}
    else:
        return data_init.update({name_of_point:file_data})

def get_single_file_name(filename):
    p1 = filename.rfind('\\')
    p2 = filename.rfind('.')
    return filename[p1+1:p2]

def statistics(data,angle1,ang,zhang_or_luo = False):##zhang_or_luo为True时，angle为落潮流向,data为某一点报表格式,ang转流包括角度

    def is_time_of(d):
        def raise_or_not(angle1,ang):  # 涨/落潮主方向
            if angle1 > 180:
                raise ValueError
                return "Wrong with the angle"
            if angle1 < 90:
                angle1 = angle1 + 90
            else:
                angle1 = angle1 - 90
            angle2 = angle1 + 180
            if min(abs(d - angle1), abs(d - angle2), 360 - abs(d - angle2), 360 - abs(d - angle1)) < ang:
                return 0  # 转流
            elif angle1 < d < angle2:
                return 1  # 涨/落潮
            else:
                return -1  # 落/涨潮
        return raise_or_not

    def get_extreme(v_and_d):  # 统计涨落潮极值
        extreme_v = v_and_d.v.max()

        extreme_d = v_and_d.loc[v_and_d['v'] == extreme_v]['d']

        extreme_t = v_and_d.loc[v_and_d['v'] == extreme_v]['time']

        v_and_d['x'] = v_and_d['time'].apply(lambda x: x.minute)
        v_and_d2 = v_and_d[v_and_d['x'] == 0]
        v_and_d2['v_e'] = v_and_d2.apply(lambda df: east(df['v'], df['d']), axis=1)
        v_and_d2['v_n'] = v_and_d2.apply(lambda df: north(df['v'], df['d']), axis=1)
        mean = velocity(v_and_d2.v_e.mean(), v_and_d2.v_n.mean())  # 平均流速为矢量各分量平均之后的合成流速
        mean_d = direction(v_and_d2.v_e.mean(), v_and_d2.v_n.mean())
        return dict(max_v=extreme_v, max_d=extreme_d.values[0], mean=mean, max_t=extreme_t.values[0],mean_d = mean_d)

    def get_duration(time_v_d):  # 统计涨落潮时间
        time_v_d = time_v_d[time_v_d.timeof != 0]
        time_v_d.index = range(0, len(time_v_d))
        raise_time = 0  # 此时无法区分涨落，仅作标识用途
        ebb_time = 0
        raise_duration = datetime.timedelta(0)
        ebb_duration = datetime.timedelta(0)
        raise_last = False
        ebb_last = False
        raise_times = []
        ebb_times = []

        for i in range(1, len(time_v_d)):
            if (time_v_d.loc[i, 'timeof'] == 1) and (time_v_d.loc[i - 1, 'timeof'] == 1) and raise_last:  # 涨潮last
                raise_duration += (time_v_d.loc[i, 'time'] - time_v_d.loc[i - 1, 'time'])
                continue

            if (time_v_d.loc[i, 'timeof'] == -1) and (time_v_d.loc[i - 1, 'timeof'] == -1) and ebb_last:  # 落潮last
                ebb_duration += (time_v_d.loc[i, 'time'] - time_v_d.loc[i - 1, 'time'])
                continue

            if (time_v_d.loc[i, 'timeof'] == -1) and (time_v_d.loc[i - 1, 'timeof'] == 1):  # 涨潮end,直接变落潮
                if raise_last:
                    raise_time += 0.5
                    raise_duration += (time_v_d.loc[i, 'time'] - time_v_d.loc[i - 1, 'time']) / 2
                    raise_last = False
                    raise_times.append(raise_duration)
                    raise_duration = datetime.timedelta(0)
                ebb_time += 0.5
                ebb_last = True
                ebb_duration += (time_v_d.loc[i, 'time'] - time_v_d.loc[i - 1, 'time']) / 2
                continue

            if (time_v_d.loc[i, 'timeof'] == 1) and (time_v_d.loc[i - 1, 'timeof'] == -1):  # 落潮end,直接变涨潮
                if ebb_last:
                    ebb_time += 0.5
                    ebb_duration += (time_v_d.loc[i, 'time'] - time_v_d.loc[i - 1, 'time']) / 2
                    ebb_last = False
                    ebb_times.append(ebb_duration)
                    ebb_duration = datetime.timedelta(0)
                raise_time += 0.5
                raise_last = True
                raise_duration += (time_v_d.loc[i, 'time'] - time_v_d.loc[i - 1, 'time']) / 2
                continue
        #print('raise_time=' + str(raise_time))  # 次数
        #print('ebb_time=' + str(ebb_time))  # 次数
        return {'durations_1': raise_times, 'durations_2': ebb_times, 'times1': raise_time, 'times2': ebb_time,
                'last_time1': reduce(operator.add, raise_times) / int(raise_time),
                'last_time2': reduce(operator.add, ebb_times) / int(ebb_time)}

    fenceng = dict(biaoceng = data.iloc[:,[0,1,2]],
                  ceng2 = data.iloc[:,[0,3,4]],
                  ceng4=data.iloc[:,[0,5,6]],
                  ceng6=data.iloc[:,[0,7,8]],
                  ceng8=data.iloc[:,[0,9,10]],
                  diceng = data.iloc[:,[0,11,12]],
                   ave = data.iloc[:,[0,15,16]])
    max_1 = {}
    max_2 = {}
    for i in sorted(fenceng.keys()):
        fenceng[i].columns = ['time','v','d']
        fenceng[i]['timeof'] = (fenceng[i])['d'].apply(lambda x : is_time_of(x)(angle1,ang))

        fenzhangluo = fenceng[i].groupby('timeof')
        #print(str(i)+'__________________________________________')
        for ii,jj in fenzhangluo:
            if ii == 1: #
                max_1.update({i:get_extreme(jj)})#涨潮extreme

            if ii == -1:
                max_2.update({i:get_extreme(jj)})#落潮extreme
        """print('极值1')
        print(max_1)
        print('极值2')
        print(max_2)"""

    if zhang_or_luo:
        time_of_raise = {'ave_last':get_duration(fenceng['ave'])['last_time1'],'all_last':get_duration(fenceng['ave'])['durations_1']}
        time_of_ebb ={'ave_last':get_duration(fenceng['ave'])['last_time2'],'all_last':get_duration(fenceng['ave'])['durations_2']}
        max_zhang = max_1
        max_luo = max_2
    else:
        time_of_ebb = {'ave_last':get_duration(fenceng['ave'])['last_time1'],'all_last':get_duration(fenceng['ave'])['durations_1']}
        time_of_raise ={'ave_last':get_duration(fenceng['ave'])['last_time2'],'all_last':get_duration(fenceng['ave'])['durations_2']}
        max_zhang = max_2
        max_luo = max_1

    return {'raise_time':time_of_raise,'ebb_time':time_of_ebb,'raise_extreme':max_zhang,'ebb_extreme':max_luo,'processing_data':fenceng['ave']}

def output_extreme_doc(data,doc):
    def songti(p):
        for r in p.runs:
            r.font.name = '宋体'
            r._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), '宋体')

    def add_2_extreme_rows(table, temp_data):
        r_new1 = t1.add_row()
        r_new2 = t1.add_row()

        r_new1.cells[0].paragraphs[0].add_run(text=temp_data.type)  # 涨\落潮流

        r_new1.cells[1].paragraphs[0].add_run(text=temp_data.point)
        r_new1.cells[1].merge(r_new2.cells[1])

        r_new1.cells[2].paragraphs[0].add_run(text='流速(cm/s)')
        r_new2.cells[2].paragraphs[0].add_run(text='流向(°)')

        r_new1.cells[3].paragraphs[0].add_run(text=temp_data.biaoceng.v)
        r_new2.cells[3].paragraphs[0].add_run(text=temp_data.biaoceng.d)

        r_new1.cells[4].paragraphs[0].add_run(text=temp_data.ceng2.v)
        r_new2.cells[4].paragraphs[0].add_run(text=temp_data.ceng2.d)

        r_new1.cells[5].paragraphs[0].add_run(text=temp_data.ceng4.v)
        r_new2.cells[5].paragraphs[0].add_run(text=temp_data.ceng4.d)

        r_new1.cells[6].paragraphs[0].add_run(text=temp_data.ceng6.v)
        r_new2.cells[6].paragraphs[0].add_run(text=temp_data.ceng6.d)

        r_new1.cells[7].paragraphs[0].add_run(text=temp_data.ceng8.v)
        r_new2.cells[7].paragraphs[0].add_run(text=temp_data.ceng8.d)

        r_new1.cells[8].paragraphs[0].add_run(text=temp_data.diceng.v)
        r_new2.cells[8].paragraphs[0].add_run(text=temp_data.diceng.d)

        r_new1.cells[9].paragraphs[0].add_run(text=data.ave.v)
        r_new2.cells[9].paragraphs[0].add_run(text=data.ave.d)

    h1 = doc.add_heading(text='潮流', level=2)
    h1.paragraph_format.alignment =docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    h2 = doc.add_heading(text='测验概况', level=3)

    p1 = doc.add_paragraph(text='本次测验开始时间为'+chinese_time(data.time.max())+'，至'+chinese_time(data.time.min(),False)+'结束',style=doc.styles['Body Text'])#输出测验时间

    h3 = doc.add_heading(text='潮流特征值', level=3)

    p2 = doc.add_paragraph(text='各测站各层流速最大值见下表：',style=doc.styles['Body Text'])

    t1 = doc.add_table(rows = 1,cols = 10,style ='Medium List 1 Accent 1' )
    t1.autofit = True

    t1.cell(0, 0).paragraphs[0].add_run(text='潮流')
    songti(t1.cell(0, 0).paragraphs[0])
    t1.cell(0,0).paragraphs[0].runs[0].font.bold = False

    t1.cell(0, 1).paragraphs[0].add_run(text='测站')
    songti(t1.cell(0, 1).paragraphs[0])

    for i,j in enumerate(l_ceng):
        t1.cell(0, i+3).paragraphs[0].add_run(text=j)
        songti(t1.cell(0,i+3).paragraphs[0])

    t1.cell(0,1).merge(t1.cell(0,2))
    ###每个数据加两行

    for p in doc.paragraphs:
        songti(p)

    return doc

def chinese_time(timestamp,all_time = True):
    leap_month = ''
    if lunardate.LunarDate.fromSolarDate(timestamp.year, timestamp.month, timestamp.day).isLeapMonth:
        leap_month =  '闰'
    if all_time:
        return str(timestamp.year)+'年'+str(timestamp.month)+'月'+str(timestamp.day)+'日'+'(阴历'+pycnnum.num2cn(lunardate.LunarDate.fromSolarDate(timestamp.year, timestamp.month, timestamp.day).month)+leap_month+'月'+pycnnum.num2cn(lunardate.LunarDate.fromSolarDate(timestamp.year, timestamp.month, timestamp.day).day)+')'+str(timestamp.hour)+'时'
    else:
        return str(timestamp.month)+'月'+str(timestamp.day)+'日'+'(阴历'+pycnnum.num2cn(lunardate.LunarDate.fromSolarDate(timestamp.year, timestamp.month, timestamp.day).month)+leap_month+'月'+pycnnum.num2cn(lunardate.LunarDate.fromSolarDate(timestamp.year, timestamp.month, timestamp.day).day)+')'+str(timestamp.hour)+'时'

def all_point_process(all_point_data):
    stat_data = {}
    for i in sorted(all_point_data.keys()):
        #all_point_data[i]为某点的数据
        stat_da = statistics()
        for j in sorted(all_point_data[i].keys()):
            #某点一个潮型的数据j
            stat = statistics(j)
    return stat_data

def vec(e,n):
    vec = e
    for i in e.index:
        for j in e.columns:
            vec.loc[i,j] = Vector(e.loc[i,j],n.loc[i,j],0)
    return  vec

def ang(v1,v2):#返回v1至v2旋转的角度,正为顺时针，负为逆时针
    v = Vector(0,1,0)
    x = v.angle(v1)
    y = v.angle(v2)
    a = v1.angle(v2)
    if min(y - x, y + 360 - x) > 0:
        return a
    else:
        return -a
mag = lambda v1,v2: v2.magnitude()-v1.magnitude()#返回流速差
def diff_of_cell(vec,ang_not_mag = True):#ang_not_mag为真时计算角度差，否时计算流速差
    fun = ang if (ang_not_mag) else mag

    vec_dif_of_cell = pandas.DataFrame(data=None,columns=vec.columns)
    for i in range(1,len(vec.columns)-1):
        c_name1 = vec.columns[i]
        c_name2 = vec.columns[i-1]
        try:
            vec_dif_of_cell[c_name1] = vec.apply(lambda df: fun(df[c_name2], df[c_name1]), axis=1)
        except:
            print('_______')
            print(c_name1)
            print('发生错误')
            continue
    return vec_dif_of_cell

def diff_of_time(vec,ang_not_mag = True):##TO DO 应该从有效数据倒序对应相减

    fun = ang if (ang_not_mag) else mag
    vec_dif_of_time = pandas.DataFrame(data=None,index=vec.columns)
    vec_dif_of_time[0] = None
    for i in range(1,len(vec.index)-1):
        try:
            vec_dif_of_time[i] = vec.apply(lambda df: fun(df[i-1], df[i]), axis=0)
        except:
            print('_______')
            print(i)
            print('发生错误')
            continue
    return vec_dif_of_time.transpose()

def one_point_process(file,angle,ang):
    f = pandas.ExcelFile(file)
    if len(f.sheet_names) ==1:
        print('只有一个潮型')#可以加入判断是不是半月潮
    point_data = {'data_from_file':file}

    for i in  f.sheet_names:
        print('**********'+str(i)+'开始*********')
        try:
            data = f.parse(i)
            point_data.update({i:statistics(data,angle,ang)})
        except:
            pass
        print('**********' + str(i)+'结束*********')
    return point_data

def out_put_to_excel(info_of_point_all_tides,file_name = '潮流统计信息.xlsx',file_name2 = '中间处理文件.xlsx'):#输入数据分为大中小潮的dict，对应每个潮型单独统计结果
    def init_sheet(sheet):
        sheet.write(1, 1, '涨潮极值')
        sheet.write(1, 2, '最大涨潮流速')
        sheet.write(1, 3, '最大涨潮流向')
        sheet.write(1, 4, '平均涨潮流速')
        sheet.write(1, 5, '平均涨潮流向')
        sheet.write(10, 1, '落潮极值')
        sheet.write(10, 2, '最大落潮流速')
        sheet.write(10, 3, '最大落潮流向')
        sheet.write(10, 4, '平均落潮流速')
        sheet.write(10, 5, '平均落潮流向')

        sheet.write(2, 1, '表层')
        sheet.write(3, 1, '0.2层')
        sheet.write(4, 1, '0.4层')
        sheet.write(5, 1, '0.6层')
        sheet.write(6, 1, '0.8层')
        sheet.write(7, 1, '底层')
        sheet.write(8, 1, '垂线平均')

        sheet.write(11, 1, '表层')
        sheet.write(12, 1, '0.2层')
        sheet.write(13, 1, '0.4层')
        sheet.write(14, 1, '0.6层')
        sheet.write(15, 1, '0.8层')
        sheet.write(16, 1, '底层')
        sheet.write(17, 1, '垂线平均')

        sheet.write(1, 8, '所有涨潮时间')
        sheet.write(10, 8, '所有落潮时间')

        sheet.write(1, 10, '涨潮平均时间')
        sheet.write(10, 10, '落潮平均时间')

    def data_into_sheet(data,sheet):
        r_or_e = ['raise_extreme','ebb_extreme']
        ceng = ['biaoceng','ceng2','ceng4','ceng6','ceng8','diceng','ave']
        vdt = ['max_v','max_d','mean','mean_d']
        for i in range(len(r_or_e)):
            for j in range(len(ceng)):
                for k in  range(len(vdt)):
                    try:
                        sheet.write(i*9+j+2,k+2,str(data[r_or_e[i]][ceng[j]][vdt[k]]))
                    except:
                        sheet.write(i * 9 + j + 2, k + 2, 'ERROR'+r_or_e[i]+' '+ceng[j]+' '+vdt[k])

        r_or_e2 = ['raise_time','ebb_time']
        for i in range(len(r_or_e2)):
            for j in range(len(data[r_or_e2[i]]['all_last'])):
                sheet.write(i*9+j+2,8,str(data[r_or_e2[i]]['all_last'][j]))

        sheet.write(3, 10, str(data['raise_time']['ave_last']))
        sheet.write(12, 10, str(data['ebb_time']['ave_last']))

    os.chdir(os.path.dirname(os.path.abspath(info_of_point_all_tides['data_from_file'])))

    point_name = get_single_file_name(info_of_point_all_tides['data_from_file'])
    excel_input = xlsxwriter.Workbook(point_name+file_name)
    excel2 = pandas.ExcelWriter(point_name+file_name2)

    for i in info_of_point_all_tides:
        one_tide = info_of_point_all_tides.get(i)
        if i == 'data_from_file':
            print('数据处理自源文件：'+one_tide)
            continue
        sheet = excel_input.add_worksheet(name=i)
        sheet.write(0,0,point_name)
        init_sheet(sheet)
        data_into_sheet(one_tide,sheet)

        one_tide['processing_data'].to_excel(excel2,sheet_name=i,index=False)

    excel_input.close()
    excel2.save()

def output_format_Excel_1(all_data,output_filename ="测流资料流速统计.xlsx" ):#输出对应各层流速平均值、极值
    # all_data[点位][潮型][涨落潮极值][层数][max/mean]
    #5-9更新可以写入不区分大小潮的涨落信息
    excel_input = xlsxwriter.Workbook(output_filename)
    format = excel_input.add_format({
        'border': 1,
        'font_name': 'Times New Roman',
        'font_size': 10,
        'num_format': '0',
        'valign': 'vcenter',
        'align': 'center'
    })
    def sheet_init(worksheet):
        worksheet.write(0, 0, '测站')
        for i,j in enumerate(['表层','0.2H','0.4H','0.6H','0.8H','底层','垂线平均']):
            worksheet.write(1, 2 * i + 1, '(cm/s)',format)
            worksheet.write(1, 2 * i + 2, '(°)',format)
            worksheet.write(0, 2 * i + 1, j,format)
            worksheet.write(0, 2 * i + 2, j,format)
    def wirte(i, p, t=None):
        for t_kind in ['raise_extreme', 'ebb_extreme']:
            for iii, ceng in enumerate(['biaoceng', 'ceng2', 'ceng4', 'ceng6', 'ceng8', 'diceng', 'ave']):
                for index, max_or_mean in enumerate(['流速最大值', '流速平均值']):
                    if t:
                        sheetname = trans(t + t_kind + max_or_mean)
                    else:
                        sheetname = trans(t_kind + max_or_mean)
                    try:
                        worksheet = excel_input.add_worksheet(name=sheetname)
                        sheet_init(worksheet)
                    except:
                        worksheet = excel_input.get_worksheet_by_name(sheetname)
                        worksheet.write(i + 2, 0, p, format)

                    if index:
                        v_d = ['mean', 'mean_d']
                    else:
                        v_d = ['max_v', 'max_d']
                    if t:
                        for index2, v_or_d in enumerate(v_d):
                            worksheet.write(i + 2, iii * 2 + 1 + index2, all_data[p][t][t_kind][ceng][v_or_d], format)
                    else:
                        for index2, v_or_d in enumerate(v_d):
                            worksheet.write(i + 2, iii * 2 + 1 + index2, all_data[p][t_kind][ceng][v_or_d], format)
    def trans(x):
        x = x.replace('da','大潮')
        x = x.replace('xiao','小潮')
        x = x.replace('zhong','中潮')
        x = x.replace('raise_extreme', '涨潮')
        x = x.replace('ebb_extreme', '落潮')
        return x
    for i,p in enumerate(sorted(all_data.keys())):
        try:
            for t in ['da','zhong','xiao']:
                wirte(i, p, t)
        except KeyError:
            excel_input.close()
        except:
            wirte(i, p)

def output_format_Excel_2(all_data,output_filename ="测流资料历时统计.xlsx" ):#统计输出涨落潮历时
    excel_input = xlsxwriter.Workbook(output_filename)
    sheet = excel_input.add_worksheet('历时统计')
    for i, p in enumerate(sorted(all_data.keys())):
        sheet.write(2+i,0,p)
        try:
            for ii,t in enumerate(['da', 'zhong', 'xiao']):
                for index,r_e in enumerate(['raise_time','ebb_time']):
                    sheet.write(2+i,ii*2+1+index,str(all_data[p][t][r_e]['ave_last']).replace('0 days ',''))
        except:
            for index, r_e in enumerate(['raise_time', 'ebb_time']):
                sheet.write(2 + i,  2 + 1 + index, str(all_data[p][r_e]['ave_last']).replace('0 days ', ''))
    excel_input.close()

all_data = {}
for i in [3]:
    angle=136
    k = r"E:\★★★★★三航院项目\★★★★★鼠浪湖-2017-3月\测流资料整理\半月\P3.xlsx"
    point = one_point_process(k,angle,0)
    all_data.update({('P'+str(i)):point})
    print('*****OK*****'+str(i))

output_format_Excel_2(all_data)
